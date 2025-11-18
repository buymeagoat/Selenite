"""Export service for generating transcripts in various formats.

Supports multiple export formats:
- txt: Plain text
- srt: SubRip subtitles with timestamps
- vtt: WebVTT subtitles
- json: Raw transcript data
- docx: Microsoft Word document
- md: Markdown formatted
"""

import json
from typing import Dict, Any, List
from io import BytesIO

from app.models.job import Job


class ExportService:
    """Service for exporting transcripts in various formats."""

    @staticmethod
    def export_txt(job: Job, transcript_text: str) -> bytes:
        """Export transcript as plain text.

        Args:
            job: Job with transcript data
            transcript_text: Full transcript text

        Returns:
            Plain text bytes
        """
        return transcript_text.encode("utf-8")

    @staticmethod
    def export_srt(job: Job, segments: List[Dict[str, Any]]) -> bytes:
        """Export transcript as SRT (SubRip) subtitle format.

        Args:
            job: Job with transcript data
            segments: List of transcript segments with timestamps

        Returns:
            SRT formatted bytes
        """
        srt_lines = []
        for i, segment in enumerate(segments, start=1):
            start = segment.get("start", 0.0)
            end = segment.get("end", 0.0)
            text = segment.get("text", "").strip()

            # Format timestamps as HH:MM:SS,mmm
            start_time = ExportService._format_srt_timestamp(start)
            end_time = ExportService._format_srt_timestamp(end)

            srt_lines.append(f"{i}")
            srt_lines.append(f"{start_time} --> {end_time}")
            srt_lines.append(text)
            srt_lines.append("")  # Blank line between entries

        return "\n".join(srt_lines).encode("utf-8")

    @staticmethod
    def export_vtt(job: Job, segments: List[Dict[str, Any]]) -> bytes:
        """Export transcript as WebVTT subtitle format.

        Args:
            job: Job with transcript data
            segments: List of transcript segments with timestamps

        Returns:
            WebVTT formatted bytes
        """
        vtt_lines = ["WEBVTT", "", ""]

        for segment in segments:
            start = segment.get("start", 0.0)
            end = segment.get("end", 0.0)
            text = segment.get("text", "").strip()

            # Format timestamps as HH:MM:SS.mmm
            start_time = ExportService._format_vtt_timestamp(start)
            end_time = ExportService._format_vtt_timestamp(end)

            vtt_lines.append(f"{start_time} --> {end_time}")
            vtt_lines.append(text)
            vtt_lines.append("")

        return "\n".join(vtt_lines).encode("utf-8")

    @staticmethod
    def export_json(job: Job, transcript_data: Dict[str, Any]) -> bytes:
        """Export transcript as JSON with metadata.

        Args:
            job: Job with transcript data
            transcript_data: Full transcript data including segments

        Returns:
            JSON formatted bytes
        """
        export_data = {
            "job_id": job.id,
            "filename": job.original_filename,
            "created_at": job.created_at.isoformat() if job.created_at else None,
            "duration": job.duration,
            "language": job.language_detected,
            "model": job.model_used,
            "speaker_count": job.speaker_count,
            "text": transcript_data.get("text", ""),
            "segments": transcript_data.get("segments", []),
        }

        return json.dumps(export_data, indent=2, ensure_ascii=False).encode("utf-8")

    @staticmethod
    def export_docx(job: Job, transcript_text: str, segments: List[Dict[str, Any]]) -> bytes:
        """Export transcript as Microsoft Word document.

        Args:
            job: Job with transcript data
            transcript_text: Full transcript text
            segments: List of transcript segments (optional, for enhanced formatting)

        Returns:
            DOCX file bytes
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            raise ImportError("python-docx package required for DOCX export")

        doc = Document()

        # Add title
        title = doc.add_heading(job.original_filename, level=1)
        title.runs[0].font.color.rgb = RGBColor(0, 0, 0)

        # Add metadata
        metadata = doc.add_paragraph()
        metadata.add_run("Duration: ").bold = True
        metadata.add_run(f"{job.duration:.1f}s\n" if job.duration else "N/A\n")
        metadata.add_run("Language: ").bold = True
        metadata.add_run(f"{job.language_detected or 'Unknown'}\n")
        metadata.add_run("Model: ").bold = True
        metadata.add_run(f"{job.model_used or 'N/A'}\n")

        doc.add_paragraph()  # Spacer

        # Add transcript heading
        doc.add_heading("Transcript", level=2)

        # Add transcript text (with segments if available)
        if segments:
            for segment in segments:
                start = segment.get("start", 0.0)
                text = segment.get("text", "").strip()

                p = doc.add_paragraph()
                timestamp = p.add_run(f"[{ExportService._format_timestamp(start)}] ")
                timestamp.font.size = Pt(9)
                timestamp.font.color.rgb = RGBColor(128, 128, 128)
                p.add_run(text)
        else:
            doc.add_paragraph(transcript_text)

        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def export_md(job: Job, transcript_text: str, segments: List[Dict[str, Any]]) -> bytes:
        """Export transcript as Markdown.

        Args:
            job: Job with transcript data
            transcript_text: Full transcript text
            segments: List of transcript segments (optional)

        Returns:
            Markdown formatted bytes
        """
        md_lines = [
            f"# {job.original_filename}",
            "",
            "## Metadata",
            "",
            f"- **Duration:** {job.duration:.1f}s" if job.duration else "- **Duration:** N/A",
            f"- **Language:** {job.language_detected or 'Unknown'}",
            f"- **Model:** {job.model_used or 'N/A'}",
            f"- **Speakers:** {job.speaker_count or 'N/A'}",
            "",
            "## Transcript",
            "",
        ]

        if segments:
            for segment in segments:
                start = segment.get("start", 0.0)
                text = segment.get("text", "").strip()
                timestamp = ExportService._format_timestamp(start)
                md_lines.append(f"**[{timestamp}]** {text}")
                md_lines.append("")
        else:
            md_lines.append(transcript_text)

        return "\n".join(md_lines).encode("utf-8")

    @staticmethod
    def _format_srt_timestamp(seconds: float) -> str:
        """Format seconds as SRT timestamp (HH:MM:SS,mmm).

        Args:
            seconds: Time in seconds

        Returns:
            Formatted timestamp string
        """
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"

    @staticmethod
    def _format_vtt_timestamp(seconds: float) -> str:
        """Format seconds as WebVTT timestamp (HH:MM:SS.mmm).

        Args:
            seconds: Time in seconds

        Returns:
            Formatted timestamp string
        """
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"

    @staticmethod
    def _format_timestamp(seconds: float) -> str:
        """Format seconds as readable timestamp (MM:SS).

        Args:
            seconds: Time in seconds

        Returns:
            Formatted timestamp string
        """
        minutes = int(seconds // 60)
        secs = int(seconds % 60)
        return f"{minutes:02d}:{secs:02d}"


# Global service instance
export_service = ExportService()
