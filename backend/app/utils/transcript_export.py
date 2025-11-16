"""Utilities to format and export transcripts in various formats.

These helpers generate content on-the-fly for Increment 6 without
depending on persisted segment storage. They accept a transcript "shape"
of text + segments and return bytes with a matching content type.
"""

from __future__ import annotations

import io
import json
from datetime import timedelta
from typing import List, Tuple, Dict, Any

from docx import Document


def _format_timestamp(seconds: float, *, srt: bool = False) -> str:
    """Format seconds to SRT or VTT timecode.

    SRT: HH:MM:SS,mmm
    VTT: HH:MM:SS.mmm
    """
    if seconds < 0:
        seconds = 0
    td = timedelta(seconds=float(seconds))
    total_seconds = int(td.total_seconds())
    hours = total_seconds // 3600
    minutes = (total_seconds % 3600) // 60
    secs = total_seconds % 60
    millis = int((td.total_seconds() - total_seconds) * 1000)
    if srt:
        return f"{hours:02}:{minutes:02}:{secs:02},{millis:03}"
    return f"{hours:02}:{minutes:02}:{secs:02}.{millis:03}"


def export_txt(text: str) -> Tuple[bytes, str]:
    return text.encode("utf-8"), "text/plain"


def export_md(title: str, text: str) -> Tuple[bytes, str]:
    content = f"# Transcript: {title}\n\n{text}\n"
    return content.encode("utf-8"), "text/markdown"


def export_json(payload: Dict[str, Any]) -> Tuple[bytes, str]:
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8"), "application/json"


def export_srt(segments: List[Dict[str, Any]]) -> Tuple[bytes, str]:
    lines: List[str] = []
    for idx, seg in enumerate(segments, start=1):
        start = _format_timestamp(seg.get("start", 0.0), srt=True)
        end = _format_timestamp(seg.get("end", seg.get("start", 0.0)), srt=True)
        text = seg.get("text", "").strip()
        lines.append(str(idx))
        lines.append(f"{start} --> {end}")
        lines.append(text)
        lines.append("")
    return "\n".join(lines).encode("utf-8"), "text/srt"


def export_vtt(segments: List[Dict[str, Any]]) -> Tuple[bytes, str]:
    lines: List[str] = ["WEBVTT", ""]
    for seg in segments:
        start = _format_timestamp(seg.get("start", 0.0), srt=False)
        end = _format_timestamp(seg.get("end", seg.get("start", 0.0)), srt=False)
        text = seg.get("text", "").strip()
        lines.append(f"{start} --> {end}")
        lines.append(text)
        lines.append("")
    return "\n".join(lines).encode("utf-8"), "text/vtt"


def export_docx(
    title: str, segments: List[Dict[str, Any]], meta: Dict[str, Any] | None = None
) -> Tuple[bytes, str]:
    doc = Document()
    doc.add_heading(f"Transcript: {title}", level=1)
    if meta and meta.get("language"):
        doc.add_paragraph(f"Language: {meta['language']}")
    doc.add_paragraph("")
    for seg in segments:
        text = seg.get("text", "").strip()
        if not text:
            continue
        # Include timestamps inline for basic context
        start = _format_timestamp(seg.get("start", 0.0), srt=False)
        end = _format_timestamp(seg.get("end", seg.get("start", 0.0)), srt=False)
        doc.add_paragraph(f"[{start} - {end}] {text}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
